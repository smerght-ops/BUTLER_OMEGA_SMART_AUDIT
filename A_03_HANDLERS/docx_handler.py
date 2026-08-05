import os
import shutil
import tempfile
from pathlib import Path

from A_03_HANDLERS.base_handler import BaseHandler

class DocxHandler(BaseHandler):

    supported_extensions = [".docx"]

    def extract(self, path: Path) -> dict:
        """
        Extract text content from a DOCX file.

        Returns:
            {
                "success": bool,
                "text": str,  # Extracted text
                "metadata": dict
            }
        """
        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "EXTRACT_FAILED",
                    "text": "Библиотека python-docx недоступна.", "metadata": {}}

        try:
            document = Document(path)
            text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
            return {
                "success": True,
                "text": text,
                "metadata": {
                    "path": str(path),
                    "paragraph_count": len(document.paragraphs)
                }
            }
        except Exception as exc:
            return {"success": False, "error": "EXTRACT_FAILED",
                    "text": f"Не удалось извлечь текст: {exc}",
                    "metadata": {"path": str(path)}}

    def format_document(self, target, operation, value=None):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ImportError:
            return {"success": False, "error": "FORMATTING_FAILED", "text": "Библиотека python-docx недоступна.", "metadata": {}}

        target = Path(str(target))
        if not target.is_file() or target.suffix.lower() != ".docx":
            return {"success": False, "error": "DOCUMENT_NOT_FOUND", "text": "Активный Word-документ не найден.", "metadata": {"target_path": str(target)}}

        temporary = None
        try:
            document = Document(target)
            changed_paragraphs = 0
            changed_runs = 0

            title = document.paragraphs[0] if document.paragraphs else None
            body = [paragraph for paragraph in document.paragraphs[1:] if paragraph.text.strip()]

            if operation in {"bold_title", "center_title", "title_font_size"}:
                if title is None or not title.text.strip():
                    return {"success": False, "error": "FORMAT_TARGET_NOT_FOUND", "text": "Заголовок для форматирования не найден.", "metadata": {"target_path": str(target)}}
            if operation == "justify_body" and not body:
                return {"success": False, "error": "FORMAT_TARGET_NOT_FOUND", "text": "Основной текст для форматирования не найден.", "metadata": {"target_path": str(target)}}

            if operation == "bold_title":
                for run in title.runs:
                    if run.text and run.bold is not True:
                        run.bold = True
                        changed_runs += 1
                changed_paragraphs = 1 if changed_runs else 0
            elif operation == "center_title":
                if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    changed_paragraphs = 1
            elif operation == "title_font_size":
                if value is None or not 1 <= float(value) <= 1638:
                    return {"success": False, "error": "INVALID_COMMAND", "text": "Некорректный размер шрифта.", "metadata": {"target_path": str(target)}}
                expected_size = Pt(float(value))
                for run in title.runs:
                    if run.text and run.font.size != expected_size:
                        run.font.size = expected_size
                        changed_runs += 1
                changed_paragraphs = 1 if changed_runs else 0
            elif operation == "justify_body":
                for paragraph in body:
                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        changed_paragraphs += 1
            elif operation != "save":
                return {"success": False, "error": "INVALID_COMMAND", "text": "Команда оформления документа не поддерживается.", "metadata": {"target_path": str(target)}}

            if operation != "save" and changed_paragraphs == 0 and changed_runs == 0:
                return {"success": False, "error": "NO_CHANGES_MADE", "text": "Форматирование уже применено; изменений не внесено.", "metadata": {"operation": operation, "target_path": str(target), "changed_paragraphs": 0, "changed_runs": 0}}

            fd, temporary_name = tempfile.mkstemp(prefix=f".{target.name}.", suffix=".docx", dir=str(target.parent))
            os.close(fd)
            temporary = Path(temporary_name)
            document.save(temporary)
            verified = Document(temporary)
            if operation == "bold_title" and not all(run.bold is True for run in verified.paragraphs[0].runs if run.text):
                raise ValueError("FORMAT_VERIFICATION_FAILED")
            if operation == "center_title" and verified.paragraphs[0].alignment != WD_ALIGN_PARAGRAPH.CENTER:
                raise ValueError("FORMAT_VERIFICATION_FAILED")
            if operation == "title_font_size" and not all(run.font.size == Pt(float(value)) for run in verified.paragraphs[0].runs if run.text):
                raise ValueError("FORMAT_VERIFICATION_FAILED")
            if operation == "justify_body" and not all(paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY for paragraph in verified.paragraphs[1:] if paragraph.text.strip()):
                raise ValueError("FORMAT_VERIFICATION_FAILED")
            try:
                os.replace(str(temporary), str(target))
            except PermissionError:
                # LibreOffice may keep the opened target non-replaceable while
                # still allowing its contents to be updated in place.
                shutil.copyfile(str(temporary), str(target))
                temporary.unlink()
            temporary = None
            return {
                "success": True,
                "text": "DOCX успешно сохранён." if operation == "save" else "Оформление DOCX успешно применено.",
                "metadata": {"operation": operation, "target_path": str(target), "size_bytes": target.stat().st_size, "changed_paragraphs": changed_paragraphs, "changed_runs": changed_runs},
            }
        except PermissionError as exc:
            return {"success": False, "error": "SAVE_FAILED", "text": f"Не удалось сохранить DOCX: {exc}", "metadata": {"target_path": str(target)}}
        except Exception as exc:
            code = "SAVE_FAILED"
            return {"success": False, "error": code, "text": f"Не удалось обработать DOCX: {exc}", "metadata": {"target_path": str(target)}}
        finally:
            if temporary is not None and temporary.exists():
                temporary.unlink()

    def create_from_text(self, target, text, open_after_create=False):
        """
        Create a new Word document from text content.

        Args:
            target: Path to the output .docx file
            text: Text content to write into the document
            open_after_create: Whether to mark for opening after creation

        Returns:
            {
                "success": bool,
                "text": str,  # Status message
                "metadata": dict  # Contains target_path, size_bytes, etc.
            }
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ImportError:
            return {"success": False, "error": "FORMATTING_FAILED",
                    "text": "Библиотека python-docx недоступна.", "metadata": {}}

        try:
            target = Path(str(target))

            # Create new document
            document = Document()

            # Write text content - split by lines and create paragraphs
            if text:
                for line in text.split('\n'):
                    paragraph = document.add_paragraph(line.strip())
                    # Apply justified alignment for body text (business style)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Save the document
            document.save(target)

            return {
                "success": True,
                "text": f"DOCX создан: {target.name}",
                "metadata": {
                    "target_path": str(target),
                    "size_bytes": target.stat().st_size if target.exists() else 0,
                    "open_after_create": open_after_create
                }
            }
        except PermissionError as exc:
            return {"success": False, "error": "SAVE_FAILED",
                    "text": f"Не удалось создать DOCX: {exc}",
                    "metadata": {"target_path": str(target)}}
        except Exception as exc:
            return {"success": False, "error": "CREATE_FAILED",
                    "text": f"Не удалось создать DOCX: {exc}",
                    "metadata": {"target_path": str(target)}}
