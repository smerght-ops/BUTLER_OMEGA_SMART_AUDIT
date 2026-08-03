from __future__ import annotations

import hashlib
import json
import os
import shutil
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from A_03_ORCHESTRATION.dispatcher_bridge_v2 import dispatch


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def state(folder: Path) -> dict:
    return {
        str(p.relative_to(folder)): {
            "size": p.stat().st_size,
            "sha256": digest(p),
            "mtime_ns": p.stat().st_mtime_ns,
        }
        for p in sorted(folder.rglob("*")) if p.is_file()
    }


def fixtures(folder: Path) -> dict[str, Path]:
    from docx import Document
    from openpyxl import Workbook
    from pypdf import PdfWriter

    folder.mkdir(parents=True)
    paths = {ext: folder / f"sample{ext}" for ext in (".txt", ".md", ".log", ".docx", ".csv", ".xlsx", ".pdf")}
    paths[".txt"].write_text("Контроль TXT\nLatin text\nЧисло: 42", encoding="utf-8")
    paths[".md"].write_text("# Контроль MD\n\nКириллица и Latin.\n\n- 42", encoding="utf-8")
    paths[".log"].write_text("2026-07-16 INFO Контроль LOG\nvalue=42", encoding="utf-8")
    paths[".csv"].write_text("name,value\nАльфа,10\nBeta,32\n", encoding="utf-8-sig")
    doc = Document(); doc.add_heading("Контроль DOCX", 1); doc.add_paragraph("Первый абзац: кириллица."); doc.add_paragraph("Second paragraph: Latin 42.")
    table = doc.add_table(rows=2, cols=2); table.cell(0, 0).text="Ключ"; table.cell(0, 1).text="Значение"; table.cell(1, 0).text="Ответ"; table.cell(1, 1).text="42"; doc.save(paths[".docx"])
    wb=Workbook(); ws=wb.active; ws.title="Данные"; ws.append(["name","value","formula","empty"]); ws.append(["Альфа",10,"=B2*2",None]); ws2=wb.create_sheet("Second"); ws2.append(["Latin",32]); wb.save(paths[".xlsx"])
    writer=PdfWriter(); writer.add_blank_page(width=200,height=300); writer.add_metadata({"/Title":"Butler PDF control"})
    with paths[".pdf"].open("wb") as stream: writer.write(stream)
    return paths


def main() -> int:
    stamp=datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    folder=Path(f"C:/Test/ButlerDocumentsResearch_{stamp}")
    report_path=ROOT / "A_09_TESTS" / f"documents_research_evidence_{stamp}.json"
    started=datetime.now(timezone.utc).astimezone().isoformat()
    paths=fixtures(folder); baseline=state(folder); records=[]
    templates=[
        'Открой документ "{p}"', 'Прочитай документ "{p}"', 'Извлеки текст из документа "{p}"',
        'Покажи содержимое документа "{p}"', 'Проанализируй документ "{p}"',
    ]
    try:
        for ext,path in paths.items():
            current=list(templates)
            if ext in {".csv",".xlsx"}: current += ['Покажи таблицу "{p}"','Извлеки данные из таблицы "{p}"']
            for template in current:
                query=template.format(p=path)
                before=state(folder); began=time.perf_counter(); result=dispatch(query,{}); elapsed=int((time.perf_counter()-began)*1000); after=state(folder)
                records.append({"format":ext,"query":query,"pid":os.getpid(),"started_at":started,"elapsed_ms_observed":elapsed,
                                "result":json.loads(json.dumps(result,ensure_ascii=False,default=str)),"filesystem_unchanged":before==after})
        project_query="Покажи документацию проекта Butler"
        result=dispatch(project_query,{})
        records.append({"format":"PROJECT","query":project_query,"pid":os.getpid(),"started_at":started,
                        "result":json.loads(json.dumps(result,ensure_ascii=False,default=str)),"filesystem_unchanged":state(folder)==baseline})
        final_state=state(folder)
        evidence={"project_root":str(ROOT),"butler_os":str((ROOT/'BUTLER_OS.py').resolve()),"python":sys.executable,
                  "pid":os.getpid(),"started_at":started,"test_root":str(folder),"baseline":baseline,"final_state":final_state,
                  "all_read_only":baseline==final_state and all(x["filesystem_unchanged"] for x in records),"records":records}
        report_path.write_text(json.dumps(evidence,ensure_ascii=False,indent=2),encoding="utf-8")
        print(report_path)
        print(json.dumps({"pid":os.getpid(),"records":len(records),"all_read_only":evidence["all_read_only"]},ensure_ascii=False))
    finally:
        if folder.exists(): shutil.rmtree(folder)
        print(json.dumps({"temporary_data_removed":not folder.exists(),"test_root":str(folder)},ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
