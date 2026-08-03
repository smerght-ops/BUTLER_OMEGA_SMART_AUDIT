# -*- coding: utf-8 -*-
import contextlib
import io
import json
import sys
import types
from pathlib import Path
from unittest.mock import patch

requests = types.ModuleType("requests")
requests.get = lambda *args, **kwargs: None
requests.post = lambda *args, **kwargs: None
requests.Session = object
sys.modules.setdefault("requests", requests)

import BUTLER_OS


COMMANDS = [
    'Создай Word-документ с текстом "Заголовок" и "Основной текст документа".',
    "Сделай заголовок жирным.",
    "Отцентрируй заголовок.",
    "Увеличь заголовок до 16.",
    "Выровняй основной текст по ширине.",
    "Сохрани документ.",
    "exit",
]
OUTPUT = Path(__file__).resolve().parents[1] / "A_06_WORKSPACE" / "STAGE4_OUTPUT"


def files():
    return [{"path": str(path.resolve()), "ctime_ns": path.stat().st_ctime_ns, "mtime_ns": path.stat().st_mtime_ns}
            for path in sorted(OUTPUT.glob("*.docx"))]


def main():
    before = files()
    before_paths = {item["path"] for item in before}
    contracts = []
    snapshots = []
    original_dispatch = BUTLER_OS.dispatch

    def recording_dispatch(query, context=None):
        result = original_dispatch(query, context)
        contracts.append(result)
        current = files()
        snapshots.append({"query": query, "contract": result,
                          "session_files": [item for item in current if item["path"] not in before_paths]})
        return result

    stream = io.StringIO()
    with patch("builtins.input", side_effect=COMMANDS), patch.object(BUTLER_OS, "clear_screen"), patch.object(BUTLER_OS, "dispatch", recording_dispatch):
        with contextlib.redirect_stdout(stream):
            BUTLER_OS.main()

    after = files()
    created = [item for item in after if item["path"] not in before_paths]
    target_paths = [item["contract"].get("metadata", {}).get("target_path") for item in snapshots]
    assert len(created) == 1, created
    assert len(snapshots) == 6
    assert all(item["contract"].get("ok") is True for item in snapshots), snapshots
    assert len(set(target_paths)) == 1 and target_paths[0] == created[0]["path"], target_paths
    assert all(len(item["session_files"]) == 1 for item in snapshots), snapshots

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    document = Document(created[0]["path"])
    first, second = document.paragraphs[:2]
    properties = {
        "paragraph_count": len(document.paragraphs),
        "first_text": first.text,
        "first_bold": all(run.bold is True for run in first.runs if run.text),
        "first_alignment": "CENTER" if first.alignment == WD_ALIGN_PARAGRAPH.CENTER else str(first.alignment),
        "first_sizes_pt": [run.font.size.pt if run.font.size else None for run in first.runs if run.text],
        "second_text": second.text,
        "second_alignment": "JUSTIFY" if second.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else str(second.alignment),
    }
    print(json.dumps({"before_count": len(before), "created": created, "snapshots": snapshots,
                      "properties": properties, "official_output": stream.getvalue()}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
